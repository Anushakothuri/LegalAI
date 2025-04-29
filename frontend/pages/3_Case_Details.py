import streamlit as st
import requests
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from sidebar import apply_sidebar_style 

API_BASE_URL = "http://127.0.0.1:8000"

# Apply the sidebar styling
apply_sidebar_style()

st.title("📜 Case Details")

if "cases" in st.session_state:
    cases = st.session_state["cases"]

    # ✅ Format dropdown options like "Case 1 - Supreme Court of India"
    case_options = {f"Case {i+1} - {case['court_name']}": case for i, case in enumerate(cases)}

    selected_case_title = st.selectbox("Select a case:", list(case_options.keys()))
    selected_case = case_options[selected_case_title]

    if st.button("View Case Details"):
        with st.spinner("Fetching case details..."):
            response = requests.get(f"{API_BASE_URL}/case_details/", params={"index": selected_case['index']})
            if response.status_code == 200:
                case_data = response.json()
                
                # Display Case Details (Removed Extra "Case Breakdown" Header)
                st.subheader(f"📜 {case_data['court_name']}")
                st.markdown(case_data['breakdown'])  # Directly render the breakdown without adding "Case Breakdown"

                # ✅ Create a DOCX (Word) file dynamically with improved structure
                docx_buffer = BytesIO()
                doc = Document()

                # **Add Proper Headings**
                # Court Name as Title (Centered)
                title = doc.add_heading(case_data['court_name'], level=1)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Petitioner
                p = doc.add_paragraph()
                p.add_run("Petitioner: ").bold = True
                p.add_run(case_data['petitioner'])

                # Add spacing
                doc.add_paragraph()

                # Case Breakdown Heading (Only Once)
                doc.add_heading("Case Breakdown", level=2)

                # **Format Case Breakdown Properly**
                sections = case_data['breakdown'].strip().split("\n")
                current_section = None

                for line in sections:
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith("### "):
                        # Main sections (e.g., Case Breakdown, Legal Analysis) as Heading 3
                        current_section = doc.add_heading(line.replace("### ", "").strip(), level=3)
                    elif line.startswith("- *"):
                        # Key fields (e.g., Court, Petitioner Argument Strength) as bold labels
                        label, value = line.split(":", 1)
                        p = doc.add_paragraph()
                        p.add_run(label.replace("- *", "").strip() + ": ").bold = True
                        p.add_run(value.strip())
                    elif line.startswith("- "):
                        # List items (e.g., arguments, evidence) as bullet points
                        doc.add_paragraph(line.replace("- ", "• "), style="List Bullet")
                    else:
                        # Regular text (e.g., reasoning, implications)
                        doc.add_paragraph(line.strip())

                # Add spacing at the end
                doc.add_paragraph()

                # Save the document
                doc.save(docx_buffer)
                docx_buffer.seek(0)

                # ✅ Provide download button for DOCX
                st.download_button(
                    label="📥 Download as DOCX",
                    data=docx_buffer,
                    file_name=f"Case_{selected_case['index']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error("Error fetching case details.")
else:
    st.warning("Please search for cases first.")